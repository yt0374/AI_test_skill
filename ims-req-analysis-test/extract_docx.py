import docx
import os

documents = {
    "通用WMS业务操作流程": r"C:\Users\yanta\Documents\知识库文档\通用WMS业务操作流程.docx",
    "WMS开箱模块": r"C:\Users\yanta\Documents\知识库文档\WMS开箱模块.docx",
    "WMS产品需求规划": r"C:\Users\yanta\Documents\知识库文档\WMS产品需求规划.docx",
    "面料备料全流程功能": r"C:\Users\yanta\Documents\知识库文档\面料备料全流程功能.docx",
    "裁片配套绑定全流程": r"C:\Users\yanta\Documents\知识库文档\裁片配套绑定全流程.docx",
    "通用-IMS系统操作手册V1.0.2【吊挂】": r"C:\Users\yanta\Documents\知识库文档\通用-IMS系统操作手册V1.0.2【吊挂】.docx",
    "使用手册-吊挂": r"C:\Users\yanta\Documents\知识库文档\使用手册-吊挂.docx",
}

output_dir = r"C:\Users\yanta\.claude\skills\ims-req-analysis-test\references\extracted_texts"
os.makedirs(output_dir, exist_ok=True)

for name, path in documents.items():
    print(f"Processing: {name}")
    try:
        doc = docx.Document(path)
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_texts.append(cell.text.strip())
                if row_texts:
                    lines.append(" | ".join(row_texts))

        output_path = os.path.join(output_dir, f"{name}.txt")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        print(f"  -> Saved {len(lines)} lines to {output_path}")
    except Exception as e:
        print(f"  ERROR: {e}")

print("\nDone!")
